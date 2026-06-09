#!/usr/bin/env python3
"""
Shared Drive reader for the estate-mate service account.

Reads files and comments from a Google Shared Drive, handling both native
Google files and uploaded Office files (.docx / .xlsx) and PDFs. Built for the
EstateMate contract-review run, where the Vertragswerk and Ari's Diskussionsblatt
live on a Shared Drive that the Sheets/Docs-only tooling and the personal Drive
connector cannot reach.

Uses the portals google_client (service account + domain-wide delegation,
impersonating roman@estatemate.io). All Drive calls pass the all-drives flags
(supportsAllDrives / includeItemsFromAllDrives) so Shared Drive items resolve.

Run (self-contained, pulls its own deps):
    uv run \
      --with google-auth --with google-auth-oauthlib --with google-api-python-client \
      --with python-docx --with openpyxl --with pypdf \
      python3 "/Users/romansiepelmeyer/Documents/Claude Code/portals/scripts/drive_reader.py" \
      <command> [args] --project estate-mate

Commands:
    list     --folder FOLDER_ID                 # children: id, name, mimeType, size, modifiedTime
    read     --id FILE_ID                        # extracted text (Docs/docx/pdf/txt) or rows (xlsx)
    comments --id FILE_ID                        # all comments + replies (author, content, quoted text, anchor, resolved)
    dump     --folder FOLDER_ID                  # for every child: metadata + extracted content + comments (one JSON)

Examples:
    # EstateMate contract folder -> the seven doc IDs
    ... drive_reader.py list --folder 1Xyrp2FyIn8KRm0TVOEyyPDZctGlxXS67 --project estate-mate

    # Ari's Diskussionsblatt (uploaded .xlsx) content + his comments
    ... drive_reader.py read     --id 1ytiPULjhSfX0dOYzGjuEYTzvmugfca9f --project estate-mate
    ... drive_reader.py comments --id 1ytiPULjhSfX0dOYzGjuEYTzvmugfca9f --project estate-mate
"""
import argparse
import io
import json
import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).parent))
from google_client import get_drive_service

GOOGLE_DOC = "application/vnd.google-apps.document"
GOOGLE_SHEET = "application/vnd.google-apps.spreadsheet"
GOOGLE_SLIDES = "application/vnd.google-apps.presentation"
DOCX = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
XLSX = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"

META_FIELDS = "id,name,mimeType,size,modifiedTime"


def _meta(drive, file_id):
    return drive.files().get(
        fileId=file_id, fields=META_FIELDS, supportsAllDrives=True
    ).execute()


def list_folder(drive, folder_id):
    files, token = [], None
    while True:
        resp = drive.files().list(
            q=f"'{folder_id}' in parents and trashed=false",
            fields=f"nextPageToken, files({META_FIELDS})",
            supportsAllDrives=True,
            includeItemsFromAllDrives=True,
            corpora="allDrives",
            pageSize=200,
            orderBy="folder,name",
            pageToken=token,
        ).execute()
        files.extend(resp.get("files", []))
        token = resp.get("nextPageToken")
        if not token:
            break
    return files


def _download_bytes(drive, file_id):
    from googleapiclient.http import MediaIoBaseDownload

    request = drive.files().get_media(fileId=file_id, supportsAllDrives=True)
    buf = io.BytesIO()
    downloader = MediaIoBaseDownload(buf, request)
    done = False
    while not done:
        _, done = downloader.next_chunk()
    return buf.getvalue()


def _export_text(drive, file_id, mime="text/plain"):
    data = drive.files().export(fileId=file_id, mimeType=mime).execute()
    return data.decode("utf-8", "replace") if isinstance(data, bytes) else str(data)


def _docx_text(data):
    from docx import Document

    doc = Document(io.BytesIO(data))
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            parts.append(" | ".join(cell.text for cell in row.cells))
    return "\n".join(parts)


def _xlsx_sheets(data):
    from openpyxl import load_workbook

    # read_only=False so embedded cell comments load; data_only=True for cached values.
    wb = load_workbook(io.BytesIO(data), read_only=False, data_only=True)
    sheets = []
    for ws in wb.worksheets:
        rows, comments = [], []
        for row in ws.iter_rows():
            out_row = []
            for cell in row:
                out_row.append("" if cell.value is None else str(cell.value))
                cmt = getattr(cell, "comment", None)
                if cmt is not None and (cmt.text or "").strip():
                    comments.append(
                        {"cell": cell.coordinate, "author": cmt.author, "text": cmt.text}
                    )
            rows.append(out_row)
        sheet = {"name": ws.title, "rows": rows}
        if comments:
            sheet["comments"] = comments
        sheets.append(sheet)
    return sheets


def _pdf_text(data):
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(data))
    return "\n".join((page.extract_text() or "") for page in reader.pages)


def read_file(drive, file_id, meta=None):
    meta = meta or _meta(drive, file_id)
    mt = meta["mimeType"]
    out = {"id": meta.get("id", file_id), "name": meta.get("name"), "mimeType": mt}
    try:
        if mt == GOOGLE_DOC:
            out["text"] = _export_text(drive, file_id, "text/plain")
        elif mt == GOOGLE_SHEET:
            out["text"] = _export_text(drive, file_id, "text/csv")
        elif mt == DOCX:
            out["text"] = _docx_text(_download_bytes(drive, file_id))
        elif mt == XLSX:
            out["sheets"] = _xlsx_sheets(_download_bytes(drive, file_id))
        elif mt == "application/pdf":
            out["text"] = _pdf_text(_download_bytes(drive, file_id))
        elif mt.startswith("text/"):
            out["text"] = _download_bytes(drive, file_id).decode("utf-8", "replace")
        elif mt == GOOGLE_SLIDES:
            out["text"] = _export_text(drive, file_id, "text/plain")
        else:
            out["error"] = f"Unsupported mimeType: {mt}"
    except Exception as e:  # noqa: BLE001 - per-file resilience for dump
        out["error"] = f"{type(e).__name__}: {e}"
    return out


def list_comments(drive, file_id):
    comments, token = [], None
    fields = (
        "nextPageToken, comments(id,author/displayName,content,"
        "quotedFileContent/value,anchor,resolved,createdTime,modifiedTime,"
        "replies(author/displayName,content,createdTime))"
    )
    while True:
        resp = drive.comments().list(
            fileId=file_id,
            fields=fields,
            pageSize=100,
            includeDeleted=False,
            pageToken=token,
        ).execute()
        comments.extend(resp.get("comments", []))
        token = resp.get("nextPageToken")
        if not token:
            break
    return comments


def dump_folder(drive, folder_id):
    children = list_folder(drive, folder_id)
    result = {"folder": folder_id, "count": len(children), "files": []}
    for meta in children:
        if meta["mimeType"] == "application/vnd.google-apps.folder":
            result["files"].append({**meta, "skipped": "folder"})
            continue
        entry = read_file(drive, meta["id"], meta=meta)
        try:
            entry["comments"] = list_comments(drive, meta["id"])
        except Exception as e:  # noqa: BLE001
            entry["comments_error"] = f"{type(e).__name__}: {e}"
        result["files"].append(entry)
    return result


def main():
    parser = argparse.ArgumentParser(description="Shared Drive reader (estate-mate)")
    parser.add_argument("command", choices=["list", "read", "comments", "dump"])
    parser.add_argument("--folder", help="Folder ID (list, dump)")
    parser.add_argument("--id", help="File ID (read, comments)")
    parser.add_argument("--project", default="estate-mate", help="Project (default estate-mate)")
    args = parser.parse_args()

    drive = get_drive_service(project=args.project)

    try:
        if args.command == "list":
            if not args.folder:
                parser.error("list requires --folder")
            payload = {"folder": args.folder, "files": list_folder(drive, args.folder)}
        elif args.command == "read":
            if not args.id:
                parser.error("read requires --id")
            payload = read_file(drive, args.id)
        elif args.command == "comments":
            if not args.id:
                parser.error("comments requires --id")
            payload = {"id": args.id, "comments": list_comments(drive, args.id)}
        else:  # dump
            if not args.folder:
                parser.error("dump requires --folder")
            payload = dump_folder(drive, args.folder)

        print(json.dumps(payload, indent=2, ensure_ascii=False))
    except Exception as e:  # noqa: BLE001
        print(json.dumps({"success": False, "error": f"{type(e).__name__}: {e}"}, indent=2), file=sys.stderr)
        sys.exit(1)


if __name__ == "__main__":
    main()
